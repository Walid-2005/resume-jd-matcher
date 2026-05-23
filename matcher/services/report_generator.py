"""
Report generation for resume-job analysis results.

Supports PDF, DOCX, and HTML export formats.
"""
import io
import html as html_mod
from datetime import datetime

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable,
)

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ── Helpers ──────────────────────────────────────────────────────────────

def _score_label(score: int) -> str:
    if score >= 70:
        return "High Match"
    if score >= 35:
        return "Medium Match"
    return "Low Match"


def _score_color_rgb(score: int):
    """Return (r, g, b) tuple for score-based colouring."""
    if score >= 70:
        return (22, 163, 74)      # green-600
    if score >= 35:
        return (202, 138, 4)      # yellow-600
    return (220, 38, 38)          # red-600


def _ensure_list(val) -> list:
    if isinstance(val, list):
        return val
    if isinstance(val, str):
        return [val] if val else []
    return []


# ── PDF ──────────────────────────────────────────────────────────────────

def generate_pdf(data: dict) -> bytes:
    """Generate a styled PDF report and return raw bytes."""
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        topMargin=20 * mm, bottomMargin=20 * mm,
        leftMargin=20 * mm, rightMargin=20 * mm,
    )

    styles = getSampleStyleSheet()

    # Custom styles
    title_style = ParagraphStyle(
        'ReportTitle', parent=styles['Title'],
        fontSize=22, spaceAfter=4, textColor=colors.HexColor('#1e3a5f'),
    )
    subtitle_style = ParagraphStyle(
        'ReportSubtitle', parent=styles['Normal'],
        fontSize=10, textColor=colors.grey, spaceAfter=14,
    )
    heading_style = ParagraphStyle(
        'SectionHeading', parent=styles['Heading2'],
        fontSize=14, textColor=colors.HexColor('#1e3a5f'),
        spaceBefore=16, spaceAfter=6,
    )
    body_style = ParagraphStyle(
        'BodyText', parent=styles['Normal'],
        fontSize=10, leading=14, spaceAfter=4,
    )
    bullet_style = ParagraphStyle(
        'BulletItem', parent=body_style,
        leftIndent=16, bulletIndent=6,
    )

    elements = []

    # ── Header ────────────────────────────────────────────────────────
    elements.append(Paragraph("ResuMatch AI — Analysis Report", title_style))
    ts = data.get('createdAt') or datetime.now().isoformat()
    try:
        dt = datetime.fromisoformat(ts.replace('Z', '+00:00'))
        date_str = dt.strftime('%B %d, %Y at %I:%M %p')
    except Exception:
        date_str = ts
    elements.append(Paragraph(f"Generated {date_str}", subtitle_style))
    elements.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#cbd5e1')))
    elements.append(Spacer(1, 8))

    # ── Match Score ───────────────────────────────────────────────────
    score = data.get('matchScore', 0)
    r, g, b = _score_color_rgb(score)
    score_color = colors.Color(r / 255, g / 255, b / 255)
    label = _score_label(score)

    score_data = [[
        Paragraph(f'<font size="28" color="#{r:02x}{g:02x}{b:02x}"><b>{score}%</b></font>', styles['Normal']),
        Paragraph(f'<font size="12" color="#{r:02x}{g:02x}{b:02x}"><b>{label}</b></font>', styles['Normal']),
    ]]
    score_table = Table(score_data, colWidths=[80, 300])
    score_table.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('LEFTPADDING', (0, 0), (-1, -1), 8),
    ]))
    elements.append(score_table)
    elements.append(Spacer(1, 6))

    # ── Found Skills ──────────────────────────────────────────────────
    found = _ensure_list(data.get('foundSkills', []))
    if found:
        elements.append(Paragraph("Matched Skills", heading_style))
        elements.append(Paragraph(", ".join(found), body_style))

    # ── Missing Skills ────────────────────────────────────────────────
    missing = _ensure_list(data.get('missingSkills', []))
    if missing:
        elements.append(Paragraph("Missing Skills", heading_style))
        elements.append(Paragraph(", ".join(missing), body_style))

    # ── Strengths ─────────────────────────────────────────────────────
    strengths = _ensure_list(data.get('strengths', []))
    if strengths:
        elements.append(Paragraph("Strengths", heading_style))
        for s in strengths:
            elements.append(Paragraph(s, bullet_style, bulletText='•'))

    # ── Weaknesses ────────────────────────────────────────────────────
    weaknesses = _ensure_list(data.get('weaknesses', []))
    if weaknesses:
        elements.append(Paragraph("Areas for Improvement", heading_style))
        for w in weaknesses:
            elements.append(Paragraph(w, bullet_style, bulletText='•'))

    # ── Recommendations ───────────────────────────────────────────────
    recommendations = _ensure_list(data.get('recommendations', []))
    if recommendations:
        elements.append(Paragraph("Recommendations", heading_style))
        for i, rec in enumerate(recommendations, 1):
            elements.append(Paragraph(f"{i}. {rec}", body_style))

    # ── AI Insight ────────────────────────────────────────────────────
    insight = data.get('aiInsight', '')
    if insight:
        elements.append(Paragraph("AI Insight", heading_style))
        elements.append(Paragraph(insight, body_style))

    # ── Job Description ───────────────────────────────────────────────
    job_desc = data.get('jobDescription', '')
    if job_desc:
        elements.append(Spacer(1, 10))
        elements.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor('#e2e8f0')))
        elements.append(Paragraph("Job Description", heading_style))
        elements.append(Paragraph(job_desc, body_style))

    # ── Footer ────────────────────────────────────────────────────────
    elements.append(Spacer(1, 20))
    elements.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor('#e2e8f0')))
    footer_style = ParagraphStyle('Footer', parent=styles['Normal'], fontSize=8, textColor=colors.grey)
    elements.append(Paragraph("Generated by ResuMatch AI", footer_style))

    doc.build(elements)
    return buf.getvalue()


# ── DOCX ─────────────────────────────────────────────────────────────────

def generate_docx(data: dict) -> bytes:
    """Generate a .docx report and return raw bytes."""
    doc = Document()

    # Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title
    title = doc.add_heading('ResuMatch AI — Analysis Report', level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(30, 58, 95)

    ts = data.get('createdAt') or datetime.now().isoformat()
    try:
        dt = datetime.fromisoformat(ts.replace('Z', '+00:00'))
        date_str = dt.strftime('%B %d, %Y at %I:%M %p')
    except Exception:
        date_str = ts
    p = doc.add_paragraph(f'Generated {date_str}')
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = RGBColor(120, 120, 120)

    # Match Score
    score = data.get('matchScore', 0)
    r, g, b = _score_color_rgb(score)
    label = _score_label(score)

    doc.add_heading('Match Score', level=1)
    score_para = doc.add_paragraph()
    run = score_para.add_run(f'{score}%  ')
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(r, g, b)
    run2 = score_para.add_run(label)
    run2.bold = True
    run2.font.size = Pt(14)
    run2.font.color.rgb = RGBColor(r, g, b)

    # Found Skills
    found = _ensure_list(data.get('foundSkills', []))
    if found:
        doc.add_heading('Matched Skills', level=1)
        doc.add_paragraph(', '.join(found))

    # Missing Skills
    missing = _ensure_list(data.get('missingSkills', []))
    if missing:
        doc.add_heading('Missing Skills', level=1)
        doc.add_paragraph(', '.join(missing))

    # Strengths
    strengths = _ensure_list(data.get('strengths', []))
    if strengths:
        doc.add_heading('Strengths', level=1)
        for s in strengths:
            doc.add_paragraph(s, style='List Bullet')

    # Weaknesses
    weaknesses = _ensure_list(data.get('weaknesses', []))
    if weaknesses:
        doc.add_heading('Areas for Improvement', level=1)
        for w in weaknesses:
            doc.add_paragraph(w, style='List Bullet')

    # Recommendations
    recommendations = _ensure_list(data.get('recommendations', []))
    if recommendations:
        doc.add_heading('Recommendations', level=1)
        for i, rec in enumerate(recommendations, 1):
            doc.add_paragraph(f'{i}. {rec}')

    # AI Insight
    insight = data.get('aiInsight', '')
    if insight:
        doc.add_heading('AI Insight', level=1)
        doc.add_paragraph(insight)

    # Job Description
    job_desc = data.get('jobDescription', '')
    if job_desc:
        doc.add_heading('Job Description', level=1)
        doc.add_paragraph(job_desc)

    # Footer
    doc.add_paragraph('')
    footer_para = doc.add_paragraph('Generated by ResuMatch AI')
    footer_para.runs[0].font.size = Pt(8)
    footer_para.runs[0].font.color.rgb = RGBColor(150, 150, 150)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── HTML ─────────────────────────────────────────────────────────────────

def generate_html(data: dict) -> str:
    """Generate a self-contained HTML report string."""
    score = data.get('matchScore', 0)
    r, g, b = _score_color_rgb(score)
    label = _score_label(score)
    hex_color = f'#{r:02x}{g:02x}{b:02x}'

    ts = data.get('createdAt') or datetime.now().isoformat()
    try:
        dt = datetime.fromisoformat(ts.replace('Z', '+00:00'))
        date_str = dt.strftime('%B %d, %Y at %I:%M %p')
    except Exception:
        date_str = ts

    def esc(text):
        return html_mod.escape(str(text))

    def skill_badges(skills, color):
        return ' '.join(
            f'<span style="display:inline-block;padding:3px 10px;margin:2px;'
            f'border-radius:12px;font-size:13px;background:{color[0]};'
            f'color:{color[1]};border:1px solid {color[2]}">{esc(s)}</span>'
            for s in skills
        )

    found = _ensure_list(data.get('foundSkills', []))
    missing = _ensure_list(data.get('missingSkills', []))
    strengths = _ensure_list(data.get('strengths', []))
    weaknesses = _ensure_list(data.get('weaknesses', []))
    recommendations = _ensure_list(data.get('recommendations', []))
    insight = data.get('aiInsight', '')
    job_desc = data.get('jobDescription', '')

    sections = []

    # Found Skills
    if found:
        sections.append(f'''
        <div class="section">
            <h2>Matched Skills ({len(found)})</h2>
            <div>{skill_badges(found, ('#dcfce7','#166534','#bbf7d0'))}</div>
        </div>''')

    # Missing Skills
    if missing:
        sections.append(f'''
        <div class="section">
            <h2>Missing Skills ({len(missing)})</h2>
            <div>{skill_badges(missing, ('#fef2f2','#991b1b','#fecaca'))}</div>
        </div>''')

    # Strengths
    if strengths:
        items = ''.join(f'<li>{esc(s)}</li>' for s in strengths)
        sections.append(f'''
        <div class="section">
            <h2>Strengths</h2>
            <ul>{items}</ul>
        </div>''')

    # Weaknesses
    if weaknesses:
        items = ''.join(f'<li>{esc(w)}</li>' for w in weaknesses)
        sections.append(f'''
        <div class="section">
            <h2>Areas for Improvement</h2>
            <ul>{items}</ul>
        </div>''')

    # Recommendations
    if recommendations:
        items = ''.join(f'<li>{esc(r)}</li>' for r in recommendations)
        sections.append(f'''
        <div class="section">
            <h2>Recommendations</h2>
            <ol>{items}</ol>
        </div>''')

    # AI Insight
    if insight:
        sections.append(f'''
        <div class="section insight">
            <h2>AI Insight</h2>
            <p>{esc(insight)}</p>
        </div>''')

    # Job Description
    if job_desc:
        sections.append(f'''
        <div class="section">
            <h2>Job Description</h2>
            <p style="color:#555">{esc(job_desc)}</p>
        </div>''')

    body_sections = '\n'.join(sections)

    return f'''<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>ResuMatch AI — Analysis Report</title>
<style>
  * {{ margin:0; padding:0; box-sizing:border-box; }}
  body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
         color:#1e293b; background:#f8fafc; padding:40px 20px; }}
  .container {{ max-width:800px; margin:0 auto; background:#fff;
               border-radius:12px; box-shadow:0 1px 3px rgba(0,0,0,.1); padding:40px; }}
  h1 {{ font-size:26px; color:#1e3a5f; margin-bottom:4px; }}
  .date {{ font-size:13px; color:#94a3b8; margin-bottom:20px; }}
  hr {{ border:none; border-top:1px solid #e2e8f0; margin:20px 0; }}
  .score-box {{ text-align:center; padding:24px 0; }}
  .score-value {{ font-size:56px; font-weight:700; color:{hex_color}; }}
  .score-label {{ display:inline-block; margin-top:6px; padding:4px 16px;
                 border-radius:20px; font-size:14px; font-weight:600;
                 color:{hex_color}; background:rgba({r},{g},{b},0.1); }}
  .score-bar {{ width:60%; height:8px; background:#e5e7eb; border-radius:4px;
               margin:12px auto 0; overflow:hidden; }}
  .score-fill {{ height:100%; border-radius:4px; background:{hex_color};
                width:{min(score, 100)}%; }}
  .section {{ margin-top:24px; }}
  .section h2 {{ font-size:16px; color:#1e3a5f; margin-bottom:8px; }}
  .section ul, .section ol {{ padding-left:24px; }}
  .section li {{ margin-bottom:4px; line-height:1.6; }}
  .section p {{ line-height:1.7; }}
  .insight {{ background:linear-gradient(135deg,#eff6ff,#eef2ff);
             padding:20px; border-radius:8px; border:1px solid #c7d2fe; }}
  .footer {{ margin-top:30px; text-align:center; font-size:11px; color:#94a3b8; }}
</style>
</head>
<body>
<div class="container">
  <h1>ResuMatch AI — Analysis Report</h1>
  <div class="date">Generated {esc(date_str)}</div>
  <hr>
  <div class="score-box">
    <div class="score-value">{score}%</div>
    <div class="score-label">{label}</div>
    <div class="score-bar"><div class="score-fill"></div></div>
  </div>
  {body_sections}
  <hr>
  <div class="footer">Generated by ResuMatch AI</div>
</div>
</body>
</html>'''
